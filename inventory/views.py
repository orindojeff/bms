from django.urls import reverse_lazy
from django.views.generic import CreateView, UpdateView, DeleteView, ListView,TemplateView
from .models import Order, Sales, Payment
from .forms import OrderForm, SalesForm, PaymentForm
from django.shortcuts import render
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db import models
from django.template.loader import render_to_string
import io, xlwt, csv
from django.template import Context
from decimal import Decimal
from django.utils.translation import gettext as _
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment
from docx.shared import Inches
from django.views.generic import ListView
from docx import Document
from docx.shared import Inches
from inventory.models import Order
from django.views.generic import ListView
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.template.loader import get_template
from xhtml2pdf import pisa
from openpyxl import Workbook
from django.utils import timezone
from django.conf import settings
from django.http import HttpResponse, FileResponse



# Create views for Order model
class OrderCreateView(CreateView):
    model = Order
    form_class = OrderForm
    template_name = 'order_form.html'
    success_url = reverse_lazy('inventory:order-list')

class OrderUpdateView(UpdateView):
    model = Order
    form_class = OrderForm
    template_name = 'order_update.html'
    success_url = reverse_lazy('inventory:order-list')

class OrderDeleteView(DeleteView):
    model = Order
    template_name = 'order_confirm_delete.html'
    success_url = reverse_lazy('inventory:order-list')

class OrderListView(ListView):
    model = Order
    template_name = 'order_list.html'
    paginate_by = 10


    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        object_list = context['object_list']
        paginator = Paginator(object_list, self.paginate_by)

        page = self.request.GET.get('page')
        try:
            objects = paginator.page(page)
        except PageNotAnInteger:
            objects = paginator.page(1)
        except EmptyPage:
            objects = paginator.page(paginator.num_pages)

        context['object_list'] = objects
        return context


# Create views for Sales model
class SalesCreateView(CreateView):
    model = Sales
    form_class = SalesForm
    template_name = 'sales_form.html'
    success_url = reverse_lazy('inventory:sales-list')

class SalesUpdateView(UpdateView):
    model = Sales
    form_class = SalesForm
    template_name = 'sales_update.html'
    success_url = reverse_lazy('inventory:sales-list')

class SalesDeleteView(DeleteView):
    model = Sales
    template_name = 'sales_confirm_delete.html'
    success_url = reverse_lazy('inventory:sales-list')

class SalesListView(ListView):
    model = Sales
    template_name = 'sales_list.html'
    paginate_by = 10

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        object_list = context['object_list']
        paginator = Paginator(object_list, self.paginate_by)

        page = self.request.GET.get('page')
        try:
            objects = paginator.page(page)
        except PageNotAnInteger:
            objects = paginator.page(1)
        except EmptyPage:
            objects = paginator.page(paginator.num_pages)

        context['object_list'] = objects
        return context



# Create views for Payment model
class PaymentCreateView(CreateView):
    model = Payment
    form_class = PaymentForm
    template_name = 'payment_form.html'
    success_url = reverse_lazy('inventory:payment-list')

class PaymentUpdateView(UpdateView):
    model = Payment
    form_class = PaymentForm
    template_name = 'payment_update.html    '
    success_url = reverse_lazy('inventory:payment-list')


class PaymentDeleteView(DeleteView):
    model = Payment
    template_name = 'payment_confirm_delete.html'
    success_url = reverse_lazy('inventory:payment-list')

class PaymentListView(ListView):
    model = Payment
    template_name = 'payment_list.html'
    paginate_by = 10

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        object_list = context['object_list']
        paginator = Paginator(object_list, self.paginate_by)

        page = self.request.GET.get('page')
        try:
            objects = paginator.page(page)
        except PageNotAnInteger:
            objects = paginator.page(1)
        except EmptyPage:
            objects = paginator.page(paginator.num_pages)

        context['object_list'] = objects
        return context


def index(request):
    total_sales = Sales.objects.aggregate(models.Sum('daily_sales'))['daily_sales__sum']
    total_payments = Payment.objects.aggregate(models.Sum('amount'))['amount__sum']
    total_balance = Payment.objects.aggregate(models.Sum('invoice_balance'))['invoice_balance__sum']
    if total_balance is None:
        total_balance = 0
    else:
        total_balance -= total_payments
    order_items = Order.objects.order_by('-date')[:10]
    context = {
        'total_sales': total_sales,
        'total_payments': total_payments,
        'total_balance': total_balance,
        'order_items': order_items,
    }
    return render(request, 'index.html', context)






class PaginationView(TemplateView):
    template_name = 'pagination.html'
    paginate_by = 10

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        queryset = self.get_queryset()
        paginator = Paginator(queryset, self.paginate_by)
        page_number = self.request.GET.get('page')
        page_obj = paginator.get_page(page_number)
        context['page_obj'] = page_obj
        return context

    def get_queryset(self):
        # Define the queryset to paginate
        # e.g. return MyModel.objects.all()
        raise NotImplementedError("Please define the 'get_queryset' method in your view.")


    # pdf, csv, xls, docs
class OrderCSVView(ListView):
    model = Order

    def render_to_response(self, context, **response_kwargs):
        orders = self.get_queryset()

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="orders.csv"'

        writer = csv.writer(response)
        writer.writerow(['Date', 'Item', 'Quantity', 'Cost per Item', 'Total Cost'])

        for order in orders:
            writer.writerow([order.date, order.item, order.quantity, order.cost, order.total_cost])

        return response

# class OrderPDFView(ListView):
#     model = Order
#     template_name = 'order_list.html'

#     def get(self, request, *args, **kwargs):
#         orders = self.get_queryset()

#         template_path = 'order_list_pdf.html'
#         context = {'orders': orders}

#         # Create a HttpResponse object with PDF mimetype
#         response = HttpResponse(content_type='application/pdf')
#         response['Content-Disposition'] = 'attachment; filename="orders.pdf"'

#         # Render the template as a string
#         template = get_template(template_path)
#         html = template.render(context)

#         # Create a PDF object from the HTML string
#         pdf = pisa.CreatePDF(html, dest=response)

#         # Return a HttpResponse if the PDF was created successfully
#         if not pdf.err:
#             return response

#         # Return an error message if there was an issue creating the PDF
#         return HttpResponse('An error occurred while generating the PDF')

from django.views.generic.base import TemplateView
from django.http import HttpResponse
from django.template.loader import get_template
from xhtml2pdf import pisa
from io import BytesIO
from .models import Order

class OrderPDFView(TemplateView):
    template_name = 'order_list_pdf.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['orders'] = Order.objects.all()
        return context

    def render_to_pdf(self, template_src, context_dict):
        template = get_template(template_src)
        html = template.render(context_dict)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)
        if not pdf.err:
            return HttpResponse(result.getvalue(), content_type='application/pdf')
        return None

    def get(self, request, *args, **kwargs):
        context = self.get_context_data(**kwargs)
        pdf = self.render_to_pdf(self.template_name, context)
        if pdf:
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = 'filename="orders.pdf"'
            return response
        return HttpResponse("PDF rendering error", status=400)



class OrderXLSView(ListView):
    model = Order

    def get(self, request, *args, **kwargs):
        orders = self.get_queryset()

        response = HttpResponse(content_type='application/ms-excel')
        response['Content-Disposition'] = 'attachment; filename="orders.xls"'

        wb = xlwt.Workbook(encoding='utf-8')
        ws = wb.add_sheet('Orders')

        row_num = 0
        columns = ['Date', 'Item', 'Quantity', 'Cost per Item', 'Total Cost']

        font_style = xlwt.XFStyle()
        font_style.font.bold = True

        for col_num, column_title in enumerate(columns):
            ws.write(row_num, col_num, column_title, font_style)

        for order in orders:
            row_num += 1
            row = [
                timezone.localtime(order.date).strftime('%Y-%m-%d %H:%M:%S'),
                order.item,
                order.quantity,
                str(order.cost),  # Convert Money instance to string
                str(order.total_cost()),  # Call total_cost method and convert Money instance to string
            ]
            for col_num, cell_value in enumerate(row):
                ws.write(row_num, col_num, cell_value)

        wb.save(response)
        return response


class OrderDocsView(ListView):
    model = Order
    template_name = 'order_docs.html'

    def render_to_response(self, context, **response_kwargs):
        # Create a new Word document
        document = Document()

        # Add a heading to the document
        document.add_heading('Order List', 0)

        # Add a table to the document
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Item'
        hdr_cells[2].text = 'Quantity'
        hdr_cells[3].text = 'Cost per Item'
        hdr_cells[4].text = 'Total Cost'

        # Add table data
        for order in context['object_list']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(order.date)
            row_cells[1].text = order.item
            row_cells[2].text = str(order.quantity)
            row_cells[3].text = str(order.cost)
            row_cells[4].text = str(order.total_cost())

        # Save the document to a BytesIO buffer
        document_buffer = io.BytesIO()
        document.save(document_buffer)

        # Return the document as a downloadable file
        document_buffer.seek(0)
        response = FileResponse(document_buffer, as_attachment=True, filename='order_list.docx')

        return response



class PaymentListCSVView(ListView):
    model = Payment

    def render_to_response(self, context, **response_kwargs):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="payment_list.csv"'

        writer = csv.writer(response)
        writer.writerow(['Payment ID', 'Amount', 'Date'])

        payments = context['object_list']
        for payment in payments:
            writer.writerow([payment.id, payment.amount, payment.date])

        return response



class SalesPDFView(TemplateView):
    template_name = 'sales_list_pdf.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['sales'] = Sales.objects.all()
        return context

    def render_to_pdf(self, template_src, context_dict):
        template = get_template(template_src)
        html = template.render(context_dict)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)
        if not pdf.err:
            return HttpResponse(result.getvalue(), content_type='application/pdf')
        return None

    def get(self, request, *args, **kwargs):
        context = self.get_context_data(**kwargs)
        pdf = self.render_to_pdf(self.template_name, context)
        if pdf:
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = 'filename="sales.pdf"'
            return response
        return HttpResponse("PDF rendering error", status=400)


class PaymentPDFView(TemplateView):
    template_name = 'payment_list_pdf.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['payments'] = Payment.objects.all()
        return context

    def render_to_pdf(self, template_src, context_dict):
        template = get_template(template_src)
        html = template.render(context_dict)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)
        if not pdf.err:
            return HttpResponse(result.getvalue(), content_type='application/pdf')
        return None

    def get(self, request, *args, **kwargs):
        context = self.get_context_data(**kwargs)
        pdf = self.render_to_pdf(self.template_name, context)
        if pdf:
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = 'filename="Milele Bookshop Payments.pdf"'
            return response
        return HttpResponse("PDF rendering error", status=400)


class PaymentXLSView(ListView):
    model = Payment

    def get(self, request, *args, **kwargs):
        payments = self.get_queryset()

        response = HttpResponse(content_type='application/ms-excel')
        response['Content-Disposition'] = 'attachment; filename="payments.xls"'

        wb = xlwt.Workbook(encoding='utf-8')
        ws = wb.add_sheet('Payments')

        row_num = 0
        columns = ['Date', 'Invoice Number', 'Amount', 'Balance', 'Invoice Balance']

        font_style = xlwt.XFStyle()
        font_style.font.bold = True

        for col_num, column_title in enumerate(columns):
            ws.write(row_num, col_num, column_title, font_style)

        for payment in payments:
            row_num += 1
            row = [
                timezone.localtime(payment.date).strftime('%Y-%m-%d %H:%M:%S'),
                payment.invoice_number,
                str(payment.amount),
                str(payment.get_balance()),
                str(payment.invoice_balance),
            ]
            for col_num, cell_value in enumerate(row):
                ws.write(row_num, col_num, cell_value)

        wb.save(response)
        return response


class PaymentdocxView(TemplateView):
    template_name = 'payment_list.html'

    def get(self, request, *args, **kwargs):
        # Query your payment model for the necessary data
        payments = Payment.objects.all()

        # Create a new Word document
        document = Document()

        # Add a heading to the document
        document.add_heading('Payment List', 0)

        # Add a table to the document and set its column headers
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Payment Date'
        hdr_cells[1].text = 'Amount'
        hdr_cells[2].text = 'Invoice Balance'
        hdr_cells[3].text = 'Balance'

        # Add rows to the table with payment data
        for payment in payments:
            row_cells = table.add_row().cells
            row_cells[0].text = payment.date.strftime('%m/%d/%Y')
            row_cells[1].text = str(payment.amount)
            row_cells[2].text = str(payment.invoice_balance)
            row_cells[3].text = str(payment.get_balance)

        # Create a response object with the appropriate headers for a .docx file
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=payment_list.docx'

        # Save the document to the response object and return it
        document.save(response)
        return response




class SalesListPDFView(TemplateView):
    template_name = 'sales_list.html'

    def get(self, request, *args, **kwargs):
        sales = Sales.objects.all()

        context = {'sales': sales}

        # Render the template
        html_string = render_to_string(self.get_template_names(), context=context)

        # Create a PDF file
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="sales_list.pdf"'

        from xhtml2pdf import pisa
        pdf = io.BytesIO()

        pisa.CreatePDF(html_string, pdf)
        response.write(pdf.getvalue())

        return response


class SalesXlsView(ListView):
    model = Sales

    def get(self, request, *args, **kwargs):
        sales = self.get_queryset()

        response = HttpResponse(content_type='application/ms-excel')
        response['Content-Disposition'] = 'attachment; filename="sales_list.xls"'

        wb = xlwt.Workbook(encoding='utf-8')
        ws = wb.add_sheet('Sales')

        row_num = 0
        columns = ['Date', 'Daily Sales', '80% of Daily Sales', '20% of Daily Sales']

        font_style = xlwt.XFStyle()
        font_style.font.bold = True

        for col_num, column_title in enumerate(columns):
            ws.write(row_num, col_num, column_title, font_style)

        for sale in sales:
            row_num += 1
            row = [
                sale.date.strftime('%Y-%m-%d %H:%M:%S'),
                str(sale.daily_sales),
                str(sale.get_80()),
                str(sale.get_20()),
            ]
            for col_num, cell_value in enumerate(row):
                ws.write(row_num, col_num, cell_value)

        wb.save(response)
        return response


# class SalePDFView(ListView):
#     model = Sale
#     template_name = 'sale_list.html'

#     def get(self, request, *args, **kwargs):
#         sales = self.get_queryset()

#         template_path = 'sale_list_pdf.html'
#         context = {'sales': sales}

#         # Create a HttpResponse object with PDF mimetype
#         response = HttpResponse(content_type='application/pdf')
#         response['Content-Disposition'] = 'attachment; filename="sales.pdf"'

#         # Render the template as a string
#         template = get_template(template_path)
#         html = template.render(context)

#         # Create a PDF object from the HTML string
#         pdf = pisa.CreatePDF(html, dest=response)

#         # Return a HttpResponse if the PDF was created successfully
#         if not pdf.err:
#             return response

#         # Return an error message if there was an issue creating the PDF
#         return HttpResponse('An error occurred while generating the PDF')


from django.http import HttpResponse
from django.views.generic import ListView
from docx import Document
import io

from .models import Sales


class SaleDocsView(ListView):
    model = Sales
    template_name = 'sales_docs.html'

    def render_to_response(self, context, **response_kwargs):
        # Create a new Word document
        document = Document()

        # Add a heading to the document
        document.add_heading('Sales List', 0)

        # Add a table to the document
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Daily Sales'
        hdr_cells[2].text = '80% of Daily Sales'
        hdr_cells[3].text = '20% of Daily Sales'

        # Add table data
        for sale in context['object_list']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(sale.date)
            row_cells[1].text = str(sale.daily_sales)
            row_cells[2].text = str(sale.get_80())
            row_cells[3].text = str(sale.get_20())

        # Save the document to a BytesIO buffer
        document_buffer = io.BytesIO()
        document.save(document_buffer)

        # Return the document as a downloadable file
        document_buffer.seek(0)
        response = HttpResponse(document_buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=sales.docx'

        return response

